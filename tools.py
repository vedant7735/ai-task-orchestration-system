"""
tools.py — File I/O tools for the orchestration agent

Tools:
  read_pdf(path)              → str
  read_excel(path)            → dict
  write_lab_word(path, data)  → str (output path)
  write_analysis_pdf(path, data) → str (output path)
"""

import os
from datetime import datetime
import subprocess
import tempfile


# ──────────────────────────────────────────────────────────
# TOOL 1 — PDF READER
# ──────────────────────────────────────────────────────────

def read_pdf(path: str) -> str:
    """
    Extract all text from a PDF file.
    Returns plain text string.
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError("Run: pip install pymupdf")

    if not os.path.exists(path):
        raise FileNotFoundError(f"PDF not found: {path}")

    doc  = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()

    print(f"[TOOL] read_pdf: {len(text)} chars from {os.path.basename(path)}")
    return text.strip()


def extract_experiments(pdf_text: str) -> list[dict]:
    """
    Split PDF text into individual experiments.
    
    Tries to detect experiment boundaries by looking for:
      - "Experiment N", "Exp N", "Program N", "Practical N"
      - Numbered sections like "1.", "2." at the start of lines
    
    Returns list of { number, title, raw_text }
    """
    import re

    # Pattern: Experiment/Exp/Program/Practical followed by number
    pattern = re.compile(
        r'(?:Experiment|Exp|Program|Practical|Exercise)\s*[:\-]?\s*(\d+)',
        re.IGNORECASE
    )

    splits  = []
    matches = list(pattern.finditer(pdf_text))

    if not matches:
        # Fallback: split by numbered lines "1." "2." etc at line start
        pattern2 = re.compile(r'^\s*(\d+)\.\s+', re.MULTILINE)
        matches  = list(pattern2.finditer(pdf_text))

    if not matches:
        # Last resort: treat whole PDF as one experiment
        return [{
            "number":   1,
            "title":    "Experiment 1",
            "raw_text": pdf_text
        }]

    for i, match in enumerate(matches):
        start = match.start()
        end   = matches[i + 1].start() if i + 1 < len(matches) else len(pdf_text)
        chunk = pdf_text[start:end].strip()

        # Try to extract title from first line
        first_line = chunk.split('\n')[0].strip()

        splits.append({
            "number":   int(match.group(1)) if match.lastindex else i + 1,
            "title":    first_line,
            "raw_text": chunk
        })

    print(f"[TOOL] extract_experiments: found {len(splits)} experiments")
    return splits


# ──────────────────────────────────────────────────────────
# TOOL 2 — EXCEL READER
# ──────────────────────────────────────────────────────────

def read_excel(path: str) -> dict:
    """
    Read an Excel file and return structured data.
    
    Returns:
    {
        "sheets": {
            "Sheet1": {
                "headers": [...],
                "rows":    [[...], ...],
                "summary": { col: { min, max, mean, count } }
            }
        },
        "total_sheets": int,
        "total_rows":   int
    }
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError("Run: pip install openpyxl")

    if not os.path.exists(path):
        raise FileNotFoundError(f"Excel not found: {path}")

    wb     = openpyxl.load_workbook(path, data_only=True)
    result = {"sheets": {}, "total_sheets": len(wb.sheetnames), "total_rows": 0}

    for sheet_name in wb.sheetnames:
        ws      = wb[sheet_name]
        rows    = list(ws.iter_rows(values_only=True))

        if not rows:
            continue

        headers = [str(h) if h is not None else f"Col{i}" for i, h in enumerate(rows[0])]
        data    = [list(row) for row in rows[1:] if any(cell is not None for cell in row)]

        # Basic numeric summary per column
        summary = {}
        for col_idx, header in enumerate(headers):
            values = []
            for row in data:
                val = row[col_idx] if col_idx < len(row) else None
                if isinstance(val, (int, float)):
                    values.append(val)

            if values:
                summary[header] = {
                    "count": len(values),
                    "min":   round(min(values), 4),
                    "max":   round(max(values), 4),
                    "mean":  round(sum(values) / len(values), 4),
                }

        result["sheets"][sheet_name] = {
            "headers": headers,
            "rows":    data[:100],   # cap at 100 rows for context window
            "summary": summary,
        }
        result["total_rows"] += len(data)

    print(f"[TOOL] read_excel: {result['total_sheets']} sheets, {result['total_rows']} rows")
    return result


def excel_to_text(excel_data: dict) -> str:
    """Convert excel_data dict to readable text for the LLM."""
    lines = []
    for sheet_name, sheet in excel_data["sheets"].items():
        lines.append(f"=== Sheet: {sheet_name} ===")
        lines.append(f"Columns: {', '.join(sheet['headers'])}")
        lines.append(f"Rows: {len(sheet['rows'])}")

        if sheet["summary"]:
            lines.append("\nNumeric Summary:")
            for col, stats in sheet["summary"].items():
                lines.append(
                    f"  {col}: count={stats['count']}, "
                    f"min={stats['min']}, max={stats['max']}, mean={stats['mean']}"
                )

        # Show first 5 rows as sample
        lines.append("\nSample Data (first 5 rows):")
        lines.append(" | ".join(sheet["headers"]))
        lines.append("-" * 40)
        for row in sheet["rows"][:5]:
            lines.append(" | ".join(str(v) if v is not None else "" for v in row))

        lines.append("")

    return "\n".join(lines)


# ──────────────────────────────────────────────────────────
# TOOL 3 — WORD DOC WRITER (Lab File)
# ──────────────────────────────────────────────────────────

def write_lab_word(output_path: str, experiments: list[dict]) -> str:
    """
    Write a structured lab file Word document.

    experiments: list of {
        number:     int
        title:      str
        aim:        str
        theory:     str
        code:       str
        output:     str
        conclusion: str
    }

    Returns the output path.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    doc = Document()

    # ── Document title ────────────────────────────────────
    title_para = doc.add_heading("Laboratory File", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Table of contents placeholder ─────────────────────
    doc.add_heading("Table of Contents", level=1)
    for exp in experiments:
        doc.add_paragraph(
            f"Experiment {exp.get('number', '?')} — {exp.get('title', '')}",
            style="List Number"
        )
    doc.add_page_break()

    # ── Each experiment ───────────────────────────────────
    for exp in experiments:
        num   = exp.get("number", "?")
        title = exp.get("title", f"Experiment {num}")

        doc.add_heading(f"Experiment {num}", level=1)
        doc.add_heading(title, level=2)

        # Aim
        doc.add_heading("Aim", level=3)
        doc.add_paragraph(exp.get("aim", "—"))

        # Theory
        doc.add_heading("Theory", level=3)
        doc.add_paragraph(exp.get("theory", "—"))

        # Code
        doc.add_heading("Code", level=3)
        code_text = exp.get("code", "—")
        code_para = doc.add_paragraph()
        code_run  = code_para.add_run(code_text)
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(9)

        # Output
        doc.add_heading("Output", level=3)
        out_text = exp.get("output", "—")
        out_para = doc.add_paragraph()
        out_run  = out_para.add_run(out_text)
        out_run.font.name = "Courier New"
        out_run.font.size = Pt(9)

        # Conclusion
        doc.add_heading("Conclusion", level=3)
        doc.add_paragraph(exp.get("conclusion", "—"))

        doc.add_page_break()

    doc.save(output_path)
    print(f"[TOOL] write_lab_word: saved to {output_path}")
    return output_path


# ──────────────────────────────────────────────────────────
# TOOL 4 — PDF REPORT WRITER (Excel Analysis)
# ──────────────────────────────────────────────────────────

def write_analysis_pdf(output_path: str, analysis: dict) -> str:
    """
    Write a PDF analysis report.

    analysis: {
        title:      str
        summary:    str
        sections:   [ { heading: str, content: str } ]
        generated:  str (optional)
    }

    Returns the output path.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, PageBreak
        )
    except ImportError:
        raise ImportError("Run: pip install reportlab")

    doc    = SimpleDocTemplate(output_path, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading1"],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor("#0f52ba"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
        leading=16,
    )

    # Title
    story.append(Paragraph(analysis.get("title", "Analysis Report"), title_style))
    story.append(Paragraph(
        f"Generated: {analysis.get('generated', datetime.now().strftime('%d %B %Y %H:%M'))}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.5*cm))

    # Executive summary
    if analysis.get("summary"):
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(analysis["summary"], body_style))
        story.append(Spacer(1, 0.3*cm))

    # Sections
    for section in analysis.get("sections", []):
        story.append(Paragraph(section.get("heading", ""), heading_style))

        content = section.get("content", "")
        # Split on newlines and render each as a paragraph
        for line in content.split("\n"):
            line = line.strip()
            if line:
                # Escape XML special chars for reportlab
                line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(line, body_style))

        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    print(f"[TOOL] write_analysis_pdf: saved to {output_path}")
    return output_path

# ──────────────────────────────────────────────────────────
# TOOL 5 — RUN PYTHON SCRIPT
# ──────────────────────────────────────────────────────────

def run_python_code(code: str, timeout: int = 10) -> str:
    """
    Execute Python code in a temporary sandbox.
    Returns stdout + stderr.
    """
    # Basic security: don't run if code is empty or too long
    if not code or len(code) > 5000:
        return "[ERROR: Code too long or empty]"

    # Create a temp file
    with tempfile.NamedTemporaryFile(
        mode='w', suffix='.py', delete=False, encoding='utf-8'
    ) as f:
        f.write(code)
        tmp_file = f.name

    try:
        # Run it
        result = subprocess.run(
            ['python', tmp_file],
            capture_output=True,
            text=True,
            timeout=timeout,
            cwd=os.path.dirname(tmp_file)  # Run in same dir
        )
        
        stdout = result.stdout.strip()
        stderr = result.stderr.strip()
        
        output = stdout
        if stderr:
            output += f"\n\n[ERRORS]\n{stderr}"
            
        return output if output else "[No output]"
            
    except subprocess.TimeoutExpired:
        return f"[ERROR: Execution timed out after {timeout}s]"
    except Exception as e:
        return f"[ERROR: {str(e)}]"
    finally:
        # Clean up temp file
        try:
            os.unlink(tmp_file)
        except:
            pass